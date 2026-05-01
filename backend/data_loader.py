import pypdf
import json
import uuid
from pathlib import Path
from typing import Dict, Any, List
import re

CASES_DIR = Path(__file__).parent / "cases"
# Ensure cases dir exists (create parents if needed)
CASES_DIR.mkdir(parents=True, exist_ok=True)


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file bytes (with OCR fallback for scanned documents)."""
    try:
        from io import BytesIO
        # First try standard PDF extraction
        pdf_reader = pypdf.PdfReader(BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        
        # If standard extraction fails or returns very little, try OCR
        if not text or len(text.strip()) < 100:
            try:
                from .ocr_rag import extract_text_from_pdf_with_ocr
                ocr_text = extract_text_from_pdf_with_ocr(file_content)
                text = text + "\n[OCR Extracted Content]\n" + ocr_text if text else ocr_text
            except Exception as ocr_error:
                print(f"OCR not available or failed: {ocr_error}. Using standard extraction.")
        
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error extracting PDF text: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from Word document (.docx) file bytes."""
    try:
        from docx import Document
        from io import BytesIO
        
        doc = Document(BytesIO(file_content))
        text = ""
        
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        return text.strip() if text.strip() else ""
    except Exception as e:
        raise ValueError(f"Error extracting Word document text: {str(e)}")


def extract_text_from_image(file_content: bytes) -> str:
    """Extract text from image file using OCR (PIL + Tesseract)."""
    try:
        from PIL import Image
        import pytesseract
        from io import BytesIO
        
        img = Image.open(BytesIO(file_content))
        # Convert to RGB if necessary (for PNG with transparency, grayscale, etc.)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Extract text using Tesseract OCR
        text = pytesseract.image_to_string(img)
        return text.strip() if text.strip() else ""
    except Exception as e:
        raise ValueError(f"Error extracting text from image using OCR: {str(e)}")


def clean_text(text: str) -> str:
    """Clean and normalize text."""
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep punctuation
    text = re.sub(r'[^\w\s.,;:!?\'\"()-]', '', text)
    return text.strip()


def save_case(case_data: Dict[str, Any]) -> str:
    """Save case data to JSON file and return case_id."""
    case_id = str(uuid.uuid4())
    case_data['case_id'] = case_id
    
    case_file = CASES_DIR / f"{case_id}.json"
    with open(case_file, 'w', encoding='utf-8') as f:
        json.dump(case_data, f, indent=2, ensure_ascii=False)
    
    return case_id


def load_case(case_id: str) -> Dict[str, Any]:
    """Load case data from JSON file."""
    case_file = CASES_DIR / f"{case_id}.json"
    if not case_file.exists():
        raise FileNotFoundError(f"Case {case_id} not found")
    
    with open(case_file, 'r', encoding='utf-8') as f:
        return json.load(f)


def update_case(case_id: str, updates: Dict[str, Any]) -> None:
    """Update case data with new information."""
    case_data = load_case(case_id)
    case_data.update(updates)
    
    case_file = CASES_DIR / f"{case_id}.json"
    with open(case_file, 'w', encoding='utf-8') as f:
        json.dump(case_data, f, indent=2, ensure_ascii=False)
